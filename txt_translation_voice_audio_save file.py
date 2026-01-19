import streamlit as st
from deep_translator import GoogleTranslator
from gtts import gTTS
import io
import base64

from docx import Document
from PyPDF2 import PdfReader
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

from docx.shared import Pt
from docx.oxml.ns import qn

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# -----------------------------
# APP TITLE
# -----------------------------

st.set_page_config(page_title="Text Translation & Voice Generator")
st.title("Text Translation & Voice Generation")

# -----------------------------
# LANGUAGE MAP
# -----------------------------

LANG_MAP = {
    "tamil": "ta",
    "hindi": "hi",
    "telugu": "te",
    "malayalam": "ml",
    "kannada": "kn",
    "marathi": "mr",
    "gujarati": "gu",
    "bengali": "bn",
    "urdu": "ur",
    "english": "en"
}

target_language = st.selectbox(
    "Select the target language:",
    list(LANG_MAP.keys())
)

# -----------------------------
# FILE UPLOAD
# -----------------------------

uploaded_file = st.file_uploader(
    "Upload your Master file",
    type=["txt", "docx", "pdf"]
)

# -----------------------------
# SESSION STATE
# -----------------------------

if "translated_text" not in st.session_state:
    st.session_state.translated_text = ""

# -----------------------------
# FILE TEXT EXTRACTION
# -----------------------------

def extract_text(file):

    file_type = file.name.split(".")[-1].lower()
    if file_type == "txt":
        return file.read().decode("utf-8")

    elif file_type == "docx":
        doc = Document(file)
        return "\n".join(
            p.text for p in doc.paragraphs if p.text.strip()
        )

    elif file_type == "pdf":
        reader = PdfReader(file)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)

    return ""

# -----------------------------
# CHUNKED TRANSLATION
# -----------------------------

def split_text_into_chunks(text, max_chars=4500):
    
    chunks = []
    start = 0
    length = len(text)

    while start < length:
        
        end = start + max_chars
        chunk = text[start:end]

        if end < length:
            
            last_space = chunk.rfind(" ")
            
            if last_space != -1:
                end = start + last_space
                chunk = text[start:end]

        chunks.append(chunk.strip())
        start = end

    return chunks


def chunked_translate(text, target_lang):

    translator = GoogleTranslator(
        source="auto",
        target=target_lang
    )

    chunks = split_text_into_chunks(text)
    translated_chunks = []

    progress = st.progress(0)

    for i, chunk in enumerate(chunks):
        translated_chunks.append(translator.translate(chunk))
        progress.progress((i + 1) / len(chunks))

    return "\n\n".join(translated_chunks)

# -----------------------------
# EXPORT HELPERS
# -----------------------------

def generate_docx(text):
    
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Nirmala UI'
    font.size = Pt(11)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)  # ✅ DEFINE run first
    run.font.name = 'Nirmala UI'
    #run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')

    for line in text.split("\n"):
        p = doc.add_paragraph(line)
        #p._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')


    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer




def generate_pdf(text):
    
    buffer = io.BytesIO()

    # Register Unicode font (supports Indic + Urdu)
    
    pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))

    styles = getSampleStyleSheet()
    styles['Normal'].fontName = 'HeiseiMin-W3'
    styles['Normal'].fontSize = 10

    doc = SimpleDocTemplate(buffer)
    story = []

    for line in text.split("\n"):
        story.append(Paragraph(line, styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer 

# -----------------------------
# MAIN LOGIC
# -----------------------------

if uploaded_file:

    try:
        english_content = extract_text(uploaded_file)

        st.subheader("Original Text")
        st.text_area("Input", english_content, height=250)

        if st.button("Translate Now"):
            
            st.session_state.translated_text = chunked_translate(
                english_content,
                LANG_MAP[target_language]
            )

    except Exception as e:
        st.error(f"File processing failed: {e}")

# -----------------------------
# TRANSLATED OUTPUT
# -----------------------------

if st.session_state.translated_text:
    st.subheader(f"Translated Text ({target_language.capitalize()})")

    st.text_area(
        "Result",
        st.session_state.translated_text,
        height=300
    )

    # TXT download

    st.download_button(
        "Download TXT",
        st.session_state.translated_text,
        file_name=f"translated_{target_language}.txt",
        mime="text/plain"
    )

    # DOCX download

    st.download_button(
        "Download Word (.docx)",
        generate_docx(st.session_state.translated_text),
        file_name=f"translated_{target_language}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # PDF download

    st.download_button(
        "Download PDF",
        generate_pdf(st.session_state.translated_text),
        file_name=f"translated_{target_language}.pdf",
        mime="application/pdf"
    )


    # st.code(st.session_state.translated_text)

    # -----------------------------
    # TEXT TO SPEECH
    # -----------------------------

    if st.button("Generate Voice Output"):

        try:

            tts = gTTS(
                text=st.session_state.translated_text,
                lang=LANG_MAP[target_language],
                slow=False
            )

            filename = uploaded_file.name.rsplit(".", 1)[0]
            
            audio_buffer = io.BytesIO()
            tts.write_to_fp(audio_buffer)
            audio_buffer.seek(0)

            st.audio(audio_buffer, format="audio/mp3")

            st.download_button(
                "Download Audio",
                audio_buffer,
                file_name= filename + "_" + target_language + ".mp3",
                mime="audio/mpeg"
            )

            st.success("Audio generated successfully.")

        except Exception as e:
            st.error(f"Audio generation failed: {e}")
